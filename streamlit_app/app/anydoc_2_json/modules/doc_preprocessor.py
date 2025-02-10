# app/anydoc_2_json/modules/doc_preprocessor.py

import os
import re
import datetime
from datetime import timedelta
from dateutil.parser import parse as parse_date, ParserError
from typing import Optional, Union, List, Dict
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import traceback

try:
    import docx  # python-docx for DOCX file manipulation
except ImportError:
    # Handle error or include in requirements.txt
    pass

try:
    import PyPDF2  # For PDF file manipulation
except ImportError:
    pass

try:
    import datefinder  # For finding dates in text
except ImportError:
    # Handle error or include in requirements.txt
    pass

from .param_manager import ParamManager
from .logger_manager import LoggerManager

class DocPreprocessor:
    """
    DocPreprocessor class for processing documents (DOCX and PDF) according to specified
    pre-processing steps, including replacing form controls, removing texts between markers,
    find and replace text, anonymizing texts, and adjusting dates.

    Dependencies:
        - `docx` module (`python-docx` library) for DOCX file manipulation.
        - `PyPDF2` module for PDF file manipulation (limited support).
        - `re` module for regular expressions.
        - `datetime` module for date manipulations.
        - `dateutil.parser` for parsing dates in text.
        - `datefinder` module for finding dates in text.
        - `ParamManager` class for managing parameters.
        - `LoggerManager` class for logging.

    Upstream functions:
        - Any module or function that needs to preprocess documents before further processing.

    Downstream functions:
        - Preprocessing methods:
            - `replace_form_controls`
            - `remove_texts_between_markers`
            - `find_and_replace_text`
            - `anonymize_texts`
            - `adjust_dates`
        - `save_document`
    """

    def __init__(self, input_file: str, output_folder: str, param_manager: ParamManager, logger_manager: LoggerManager):
        """
        Initialize the DocPreprocessor instance.

        Parameters:
            input_file (str):
                The full path to the input document file (DOCX or PDF).
            output_folder (str):
                The directory where the processed document will be saved.
            param_manager (ParamManager):
                An instance of ParamManager containing processing parameters.
            logger_manager (LoggerManager):
                An instance of LoggerManager for logging events and errors.

        Returns:
            None

        Raises:
            FileNotFoundError:
                If the input file does not exist.
            ValueError:
                If the input file is not a DOCX or PDF file.
            IOError:
                If there is an error reading the input file.
            Exception:
                For other exceptions during initialization.

        Dependencies:
            - `os.path.exists` to check if the input file exists.
            - `docx.Document` to load DOCX files.
            - `PyPDF2.PdfReader` to load PDF files.
            - `self.logger_manager` for logging.

        Upstream functions:
            - Called by modules or functions that need to preprocess documents.

        Downstream functions:
            - `None` directly, but `process_document` and other methods are called after initialization.
        """
        self.input_file = input_file
        self.output_folder = output_folder
        self.param_manager = param_manager
        self.logger = logger_manager.get_logger()
        self.document = None
        self.doc_type = None  # 'docx' or 'pdf'

        if not os.path.exists(self.input_file):
            self.logger.error(f"Input file '{self.input_file}' does not exist.")
            raise FileNotFoundError(f"Input file '{self.input_file}' does not exist.")

        if self.input_file.lower().endswith('.docx'):
            self.doc_type = 'docx'
            try:
                self.document = docx.Document(self.input_file)
            except Exception as e:
                self.logger.error(f"Error reading DOCX file '{self.input_file}': {e}")
                raise IOError(f"Error reading DOCX file '{self.input_file}': {e}")
        elif self.input_file.lower().endswith('.pdf'):
            self.doc_type = 'pdf'
            try:
                with open(self.input_file, 'rb') as f:
                    self.document = PyPDF2.PdfReader(f)
            except Exception as e:
                self.logger.error(f"Error reading PDF file '{self.input_file}': {e}")
                raise IOError(f"Error reading PDF file '{self.input_file}': {e}")
        else:
            self.logger.error(f"Unsupported file type for '{self.input_file}'. Only DOCX and PDF are supported.")
            raise ValueError(f"Unsupported file type for '{self.input_file}'. Only DOCX and PDF are supported.")

    def replace_text_in_runs(self, element, find_text: str, replace_text: str):
        """
        Recursively replace all occurrences of a text string within runs in the document element.

        Parameters:
            element:
                - Type: docx.Document | docx.text.paragraph.Paragraph | docx.table._Cell
                - The starting element to process (can be the document itself, a table cell, or any element that contains paragraphs and tables).
            find_text (str):
                The text string to search for in runs.
            replace_text (str):
                The text string to replace the found text with.

        Returns:
            None

        Raises:
            None

        Upstream functions:
            - Called by `find_and_replace_text`.

        Downstream functions:
            - Calls `_replace_text_in_paragraph`.
            - Recursively calls itself for cells within tables.

        Dependencies:
            - Requires the `docx` module.
            - The element must be an object that has `paragraphs` and possibly `tables` attributes.

        """
        for paragraph in element.paragraphs:
            self._replace_text_in_paragraph(paragraph, find_text, replace_text)
        for table in getattr(element, 'tables', []):
            for row in table.rows:
                for cell in row.cells:
                    self.replace_text_in_runs(cell, find_text, replace_text)

    def _replace_text_in_paragraph(self, paragraph, find_text: str, replace_text: str):
        """
        Replace text in runs within a paragraph.

        Parameters:
            paragraph (docx.text.paragraph.Paragraph):
                The paragraph whose runs will be processed.
            find_text (str):
                The text string to search for in runs.
            replace_text (str):
                The text string to replace the found text with.

        Returns:
            None

        Raises:
            None

        Upstream functions:
            - Called by `replace_text_in_runs`.

        Downstream functions:
            - None

        Dependencies:
            - Requires the `docx` module.

        """
        for run in paragraph.runs:
            if find_text in run.text:
                run.text = run.text.replace(find_text, replace_text)
    
    def remove_text_between_markers_in_runs(self, element, start_marker: str, end_marker: str):
        """
        Recursively remove text between start and end markers within runs in the document element.

        Parameters:
            element:
                - Type: docx.Document | docx.text.paragraph.Paragraph | docx.table._Cell
                - The starting element to process (can be the document itself, a table cell, or any element that contains paragraphs and tables).
            start_marker (str):
                The marker indicating the start of text to be removed.
            end_marker (str):
                The marker indicating the end of text to be removed.

        Returns:
            None

        Raises:
            None

        Upstream functions:
            - Called by `remove_texts_between_markers`.

        Downstream functions:
            - Calls `_remove_text_between_markers_in_paragraph`.
            - Recursively calls itself for cells within tables.

        Dependencies:
            - Requires the `docx` module.
            - Uses the `re` module for regular expressions.

        """
        for paragraph in element.paragraphs:
            self._remove_text_between_markers_in_paragraph(paragraph, start_marker, end_marker)
        for table in getattr(element, 'tables', []):
            for row in table.rows:
                for cell in row.cells:
                    self.remove_text_between_markers_in_runs(cell, start_marker, end_marker)

    def _remove_text_between_markers_in_paragraph(self, paragraph, start_marker: str, end_marker: str):
        """
        Remove text between markers within runs of a paragraph.

        Parameters:
            paragraph (docx.text.paragraph.Paragraph):
                The paragraph whose runs will be processed.
            start_marker (str):
                The marker indicating the start of text to be removed.
            end_marker (str):
                The marker indicating the end of text to be removed.

        Returns:
            None

        Raises:
            None

        Upstream functions:
            - Called by `remove_text_between_markers_in_runs`.

        Downstream functions:
            - None

        Dependencies:
            - Requires the `docx` module.
            - Uses the `re` module for regular expressions.

        """
        # Combine all run texts
        full_text = ''.join(run.text for run in paragraph.runs)
        pattern = re.escape(start_marker) + '(.*?)' + re.escape(end_marker)
        regex = re.compile(pattern, flags=re.DOTALL)
        new_text = regex.sub(' ', full_text)
        # Clear existing runs and add new text
        paragraph.clear()
        paragraph.add_run(new_text)
    def replace_form_controls(self):
        """
        Replace interactive form controls in the document with specified text strings.

        This method now processes actual MS Word content controls (e.g., check boxes and radio buttons)
        inserted via the Developer tab in MS Word.

        Parameters:
            None

        Returns:
            None

        Raises:
            Exception:
                If there is an error during processing.

        Upstream functions:
            - Called by `process_document` if 'replaceFormControls' parameter is 'yes'.

        Downstream functions:
            - Uses methods from `docx` or `PyPDF2` to manipulate the document content.

        Dependencies:
            - `self.param_manager` for parameter values.
            - `self.logger` for logging events and errors.
            - `lxml.etree` for XML manipulation.
            - `docx.oxml.ns.qn` for XML namespaces.

        Notes:
            - Currently, only DOCX files are fully supported for this operation.
            - Handles content controls for check boxes and radio buttons.

        """
        replace_form_controls = self.param_manager.get_parameter('replaceFormControls', 'yes')
        if replace_form_controls.lower() != 'yes':
            self.logger.info("Skipping form control replacement as per parameters.")
            return

        if self.doc_type == 'docx':
            try:
                # Iterate over all paragraphs in the document
                for paragraph in self.document.paragraphs:
                    # Access the underlying XML element
                    p_element = paragraph._p
                    self._process_sdt_elements(p_element)

                # Iterate over all tables in the document
                for table in self.document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                p_element = paragraph._p
                                self._process_sdt_elements(p_element)

                self.logger.info("Form controls replaced successfully.")
            except Exception as e:
                self.logger.error(f"Error replacing form controls in DOCX file: {e}")
                raise Exception(f"Error replacing form controls in DOCX file: {e}")
        elif self.doc_type == 'pdf':
            # Modifying PDFs is non-trivial; placeholder implementation
            self.logger.warning("Form control replacement in PDFs is not fully supported.")
        else:
            self.logger.error("Unsupported document type for form control replacement.")
            raise Exception("Unsupported document type for form control replacement.")

    def _process_sdt_elements(self, parent_element):
        """
        Process all content control elements (w:sdt) in a given parent element.

        Parameters:
            parent_element: The parent XML element to search for w:sdt elements.

        Returns:
            None

        Dependencies:
            - docx.oxml.OxmlElement for creating new XML elements.
            - docx.oxml.ns.qn for qualified names.

        """

        # Find all w:sdt elements
        sdt_elements = parent_element.findall('.//' + qn('w:sdt'))
        for sdt in sdt_elements:
            sdt_pr = sdt.find(qn('w:sdtPr'))
            sdt_content = sdt.find(qn('w:sdtContent'))
            if sdt_pr is not None:
                # Check for checkbox
                check_box = sdt_pr.find(qn('w14:checkbox'))
                combo_box = sdt_pr.find(qn('w:comboBox'))
                drop_down_list = sdt_pr.find(qn('w:dropDownList'))
                # Add more content control types as needed
                if check_box is not None:
                    # Process checkbox
                    checked_elem = check_box.find(qn('w14:checked'))
                    if checked_elem is not None:
                        val_attr = checked_elem.get(qn('w14:val'))
                        checked = (val_attr == '1' or val_attr == 'true')
                    else:
                        checked = False
                    # Determine the replacement text
                    replacement_text = '{{checked box}}' if checked else '{{unchecked box}}'
                elif combo_box is not None or drop_down_list is not None:
                    # Process combo box or drop-down list
                    selected_value = ''
                    items = []
                    list_items = sdt_pr.findall('.//' + qn('w:listItem'))
                    for item in list_items:
                        val = item.get(qn('w:displayText'))
                        items.append(val)
                    # Attempt to get the selected value from sdt_content
                    if sdt_content is not None:
                        text_elems = sdt_content.findall('.//' + qn('w:t'))
                        selected_value = ''.join([t.text for t in text_elems if t.text])
                    replacement_text = f'{{{{selected value: {selected_value}}}}}'
                else:
                    # Other content controls (e.g., rich text)
                    # Get the text content
                    if sdt_content is not None:
                        text_elems = sdt_content.findall('.//' + qn('w:t'))
                        replacement_text = ''.join([t.text for t in text_elems if t.text])
                    else:
                        replacement_text = ''
                # Replace the content control with the replacement text
                # Get the parent of sdt
                parent = sdt.getparent()
                index_in_parent = parent.index(sdt)
                # Create a new run element with the replacement text
                new_run = OxmlElement('w:r')
                new_text = OxmlElement('w:t')
                new_text.text = replacement_text
                new_run.append(new_text)
                # Insert the new run into the parent and remove the sdt element
                parent.insert(index_in_parent, new_run)
                parent.remove(sdt)

    def remove_texts_between_markers(self):
        """
        Remove texts in the document that are between specified start and end markers.

        Parameters:
            None

        Returns:
            None

        Raises:
            Exception:
                If there is an error during processing.

        Upstream functions:
            - Called by `process_document`.

        Downstream functions:
            - Uses `remove_text_between_markers_in_runs`.

        Dependencies:
            - `self.param_manager` for 'removeTexts' parameter.
            - `self.logger` for logging events and errors.
            - Requires the `docx` module.

        Notes:
            - Special markers such as 'end of line' are represented by '\\n'.
            - Processes the document in-memory without altering the overall structure.

        """
        remove_texts = self.param_manager.get_parameter('removeTexts', [])
        if not remove_texts:
            self.logger.info("No text removal rules specified.")
            return

        if self.doc_type == 'docx':
            try:
                for rule in remove_texts:
                    start_marker = rule.get('start', '')
                    end_marker = rule.get('end', '')
                    if start_marker == 'end of line':
                        start_marker = '\n'
                    if end_marker == 'end of line':
                        end_marker = '\n'
                    self.logger.info(f"Removing texts between '{start_marker}' and '{end_marker}'")
                    self.remove_text_between_markers_in_runs(self.document, start_marker, end_marker)
                self.logger.info("Texts between markers removed successfully.")
            except Exception as e:
                self.logger.error(f"Error removing texts between markers: {e}")
                raise Exception(f"Error removing texts between markers: {e}")
        elif self.doc_type == 'pdf':
            self.logger.warning("Removing texts between markers in PDFs is not supported.")
        else:
            self.logger.error("Unsupported document type for removing texts between markers.")
            raise Exception("Unsupported document type for removing texts between markers.")


    def find_and_replace_text(self):
        """
        Find and replace text in the document based on specified replacement rules.

        Parameters:
            None

        Returns:
            None

        Raises:
            Exception:
                If there is an error during processing.

        Upstream functions:
            - Called by `process_document`.

        Downstream functions:
            - Uses `replace_text_in_runs`.

        Dependencies:
            - `self.param_manager` for 'replaceText' parameter.
            - `self.logger` for logging events and errors.
            - Requires the `docx` module.

        Notes:
            - Supports special characters such as 'end of line' represented by '\\n'.
            - Processes the document in-memory without altering the overall structure.

        """
        replace_texts = self.param_manager.get_parameter('replaceText', [])
        if not replace_texts:
            self.logger.info("No text replacement rules specified.")
            return

        if self.doc_type == 'docx':
            try:
                for rule in replace_texts:
                    find_text = rule.get('from', '')
                    replace_text = rule.get('to', '')
                    if not find_text:
                        continue
                    if find_text == 'end of line':
                        find_text = '\n'
                    if replace_text == 'end of line':
                        replace_text = '\n'
                    self.logger.info(f"Replacing '{find_text}' with '{replace_text}'")
                    self.replace_text_in_runs(self.document, find_text, replace_text)
                self.logger.info("Text replacement completed successfully.")
            except Exception as e:
                self.logger.error(f"Error during text replacement: {e}")
                raise Exception(f"Error during text replacement: {e}")
        elif self.doc_type == 'pdf':
            self.logger.warning("Text replacement in PDFs is not supported.")
        else:
            self.logger.error("Unsupported document type for text replacement.")
            raise Exception("Unsupported document type for text replacement.")

    def anonymize_texts(self):
        """
        Anonymize emails, person names, and organization names in the document 
        based on specified methods (redact, jibberish, realistic).

        Parameters:
            None

        Returns:
            None

        Raises:
            Exception:
                If there is an error during processing.

        Upstream functions:
            - Called by `process_document`.

        Downstream functions:
            - Uses methods from `docx` or `PyPDF2` to manipulate the document content.
            - May use external libraries for name/organization/email recognition.

        Dependencies:
            - `self.param_manager` for 'anonymization' parameter.
            - `self.logger` for logging.
            - `re` module for regex operations.
            - Possibly `spacy` or `nltk` for name entity recognition (if implemented).

        Notes:
            - Currently, only DOCX files are fully supported for this operation.
            - Basic implementation using regular expressions.
            - 'realistic' replacement may use placeholder values due to complexity.
        """

        anonymization = self.param_manager.get_parameter('anonymization', {})
        if not anonymization:
            self.logger.info("No anonymization rules specified.")
            return

        categories = anonymization.keys()
        text = ''

        if self.doc_type == 'docx':
            try:
                # Extract full text from document
                full_text = '\n'.join([para.text for para in self.document.paragraphs])
                found_entities = {'email': [], 'person name': [], 'organization': []}

                # Identify entities
                if 'email' in categories:
                    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
                    emails = re.findall(email_pattern, full_text)
                    found_entities['email'] = emails

                if 'person name' in categories or 'organization' in categories:
                    # Using placeholder for entity recognition; in a real implementation, use NLP libraries
                    # For demonstration, we'll consider any sequence of two capitalized words as names
                    if 'person name' in categories:
                        person_name_pattern = r'\b[A-Z][a-z]+\s[A-Z][a-z]+\b'
                        person_names = re.findall(person_name_pattern, full_text)
                        found_entities['person name'] = person_names
                    if 'organization' in categories:
                        # Placeholder pattern for organization names
                        organization_pattern = r'\b[A-Z][a-zA-Z0-9&]+\b(?:\s\b[A-Z][a-zA-Z0-9&]+\b)*'
                        organizations = re.findall(organization_pattern, full_text)
                        found_entities['organization'] = organizations

                self.logger.info(f"Found entities: {found_entities}")

                # Replace entities based on method
                for category, method in anonymization.items():
                    entities = set(found_entities.get(category, []))
                    self.logger.info(f"Anonymizing {len(entities)} {category}(s) using method '{method}'")
                    for entity in entities:
                        if method == 'redact':
                            replacement = '*' * len(entity)
                        elif method == 'jibberish':
                            replacement = ''.join(['*' for _ in entity])
                        elif method == 'realistic':
                            # Placeholder realistic replacements
                            if category == 'email':
                                replacement = 'anonymous@example.com'
                            elif category == 'person name':
                                replacement = 'John Doe'
                            elif category == 'organization':
                                replacement = 'Acme Corporation'
                            else:
                                replacement = '***'
                        else:
                            replacement = '***'
                        # Replace all occurrences
                        full_text = re.sub(re.escape(entity), replacement, full_text)

                # Clear existing document content and add modified text
                for para in self.document.paragraphs:
                    para.text = ''

                self.document.add_paragraph(full_text)
                self.logger.info("Anonymization completed successfully.")
            except Exception as e:
                self.logger.error(f"Error during anonymization: {e}")
                raise Exception(f"Error during anonymization: {e}")
        elif self.doc_type == 'pdf':
            # Modifying PDFs is non-trivial; placeholder implementation
            self.logger.warning("Anonymization in PDFs is not supported.")
        else:
            self.logger.error("Unsupported document type for anonymization.")
            raise Exception("Unsupported document type for anonymization.")

    def adjust_dates(self):
        """
        Adjust dates in the document based on specified rules.

        Parameters:
            None

        Returns:
            None

        Raises:
            Exception:
                If there is an error during processing.

        Upstream functions:
            - Called by `process_document`.

        Downstream functions:
            - Uses methods from `docx` or `PyPDF2` to manipulate the document content.
            - Uses `dateutil.parser` and `datetime` for date parsing and adjustments.

        Dependencies:
            - `self.param_manager` for 'adjustDates' parameter.
            - `self.logger` for logging.
            - `re` module and `datefinder` for date detection.
            - `dateutil.parser` and `datetime` for date adjustments.

        Notes:
            - Only processes dates with day, month, and year.
            - Currently, only DOCX files are fully supported for this operation.
        """

        adjust_dates = self.param_manager.get_parameter('adjustDates', {})
        if not adjust_dates:
            self.logger.info("No date adjustment rules specified.")
            return

        # Determine adjustment method
        method = None
        days = None
        for key in adjust_dates:
            if key in ['add', 'subtract', 'daysBefore', 'daysAfter']:
                method = key
                days = adjust_dates[key]
                break

        if method is None or days is None:
            self.logger.error("Invalid date adjustment parameters.")
            raise Exception("Invalid date adjustment parameters.")

        if self.doc_type == 'docx':
            try:
                for paragraph in self.document.paragraphs:
                    new_runs = []
                    text = paragraph.text
                    matches = list(datefinder.find_dates(text))
                    if matches:
                        for date_obj in matches:
                            # Only adjust dates with day, month, and year
                            if date_obj.day and date_obj.month and date_obj.year:
                                original_date_str = date_obj.strftime('%Y-%m-%d')
                                # Adjust date
                                if method == 'add':
                                    new_date = date_obj + timedelta(days=days)
                                elif method == 'subtract':
                                    new_date = date_obj - timedelta(days=days)
                                elif method == 'daysBefore':
                                    new_date = datetime.datetime.now() - timedelta(days=days)
                                elif method == 'daysAfter':
                                    new_date = datetime.datetime.now() + timedelta(days=days)
                                else:
                                    continue
                                new_date_str = new_date.strftime('%Y-%m-%d')
                                self.logger.info(f"Replacing date '{original_date_str}' with '{new_date_str}'")
                                # Replace date in text
                                text = text.replace(original_date_str, new_date_str)
                    paragraph.text = text

                self.logger.info("Date adjustment completed successfully.")
            except Exception as e:
                self.logger.error(f"Error during date adjustment: {e}")
                raise Exception(f"Error during date adjustment: {e}")
        elif self.doc_type == 'pdf':
            # Modifying PDFs is non-trivial; placeholder implementation
            self.logger.warning("Date adjustment in PDFs is not supported.")
        else:
            self.logger.error("Unsupported document type for date adjustment.")
            raise Exception("Unsupported document type for date adjustment.")

    def save_document(self):
        """
        Save the processed document to the output folder.

        Parameters:
            None

        Returns:
            None

        Raises:
            IOError:
                If there is an error saving the document.

        Dependencies:
            - Uses `self.document` to save the file.
            - `os.path.join` to construct the output file path.
            - `self.logger` for logging.

        Upstream functions:
            - Called by `process_document`.

        Downstream functions:
            - None
        """

        try:
            output_file = os.path.join(self.output_folder, os.path.basename(self.input_file))
            if self.doc_type == 'docx':
                self.document.save(output_file)
                self.logger.info(f"Document saved to '{output_file}'")
            elif self.doc_type == 'pdf':
                self.logger.warning("Saving modified PDFs is not fully supported.")
                # Placeholder; in reality, modifying and saving PDFs requires more work
            else:
                self.logger.error("Unsupported document type for saving.")
                raise Exception("Unsupported document type for saving.")
        except Exception as e:
            self.logger.error(f"Error saving document: {e}")
            raise IOError(f"Error saving document: {e}")

    def process_document(self):
        """
        Process the document by executing each pre-processing step in order.

        Parameters:
            None

        Returns:
            None

        Raises:
            Exception:
                If there is an error during processing.

        Dependencies:
            - Calls:
                - `replace_form_controls`
                - `remove_texts_between_markers`
                - `find_and_replace_text`
                - `anonymize_texts`
                - `adjust_dates`
                - `remove_empty_rows_in_tables`
                - `remove_rows_with_string_in_tables`
                - `remove_empty_columns_in_tables`
                - `save_document`
            - Uses `self.logger` for logging.

        Upstream functions:
            - Called externally to perform document pre-processing.

        Downstream functions:
            - Pre-processing methods and `save_document`.
        """

        try:
            self.replace_form_controls()
            self.remove_texts_between_markers()
            self.find_and_replace_text()
            self.anonymize_texts()
            self.adjust_dates()
            self.remove_empty_rows_in_tables()
            self.remove_rows_with_string_in_tables()
            self.remove_empty_columns_in_tables()
            self.save_document()
            self.logger.info("Document processing completed successfully.")
        except Exception as e:
            self.logger.error(f"Error during document processing: {e}")
            raise Exception(f"Error during document processing: {e}")
    
    def remove_empty_rows_in_tables(self):
        """
        Remove empty rows from all tables in the document.

        Parameters:
            None

        Returns:
            None

        Raises:
            Exception:
                If there is an error during processing.

        Upstream functions:
            - `process_document()`

        Downstream functions:
            - None

        Dependencies:
            - Requires the `docx` module.
            - The document must be a DOCX file, and `self.document` must be an instance of `docx.Document`.
            - The parameter `removeEmptyRows` must be set to `'yes'` in `self.param_manager`.
        """
        remove_empty_rows = self.param_manager.get_parameter('removeEmptyRows', 'yes')
        if remove_empty_rows.lower() != 'yes':
            self.logger.info("Skipping removing empty rows from tables as per parameters.")
            return

        if self.doc_type == 'docx':
            try:
                for table in self.document.tables:
                    rows_to_delete = []
                    for row_idx, row in enumerate(table.rows):
                        is_empty = True
                        for cell in row.cells:
                            if cell.text.strip():
                                is_empty = False
                                break
                        if is_empty:
                            rows_to_delete.append(row_idx)
                    # Delete rows from the bottom up to avoid index shifting
                    for row_idx in reversed(rows_to_delete):
                        tbl = table._tbl
                        tr = tbl.tr_lst[row_idx]
                        tbl.remove(tr)
                self.logger.info("Empty rows removed from tables successfully.")
            except Exception as e:
                self.logger.error(f"Error removing empty rows from tables: {e}")
                raise Exception(f"Error removing empty rows from tables: {e}")
        elif self.doc_type == 'pdf':
            self.logger.warning("Removing empty rows from tables in PDFs is not supported.")
        else:
            self.logger.error("Unsupported document type for removing empty rows from tables.")
            raise Exception("Unsupported document type for removing empty rows from tables.")
    
    def remove_rows_with_string_in_tables(self):
        """
        Remove rows containing a specific string from all tables in the document.

        Parameters:
            None

        Returns:
            None

        Raises:
            Exception:
                If there is an error during processing.

        Upstream functions:
            - `process_document()`

        Downstream functions:
            - None

        Dependencies:
            - Requires the `docx` module.
            - The document must be a DOCX file, and `self.document` must be an instance of `docx.Document`.
            - The parameter `removeRowsWithString` must be set in `self.param_manager`.
        """
        remove_rows_with_string = self.param_manager.get_parameter('removeRowsWithString', '')
        if not remove_rows_with_string:
            self.logger.info("No string specified for removing rows. Skipping this step.")
            return

        if self.doc_type == 'docx':
            try:
                for table in self.document.tables:
                    rows_to_delete = []
                    for row_idx, row in enumerate(table.rows):
                        row_text = ' '.join([cell.text for cell in row.cells])
                        if remove_rows_with_string in row_text:
                            rows_to_delete.append(row_idx)
                    # Delete rows from the bottom up to avoid index shifting
                    for row_idx in reversed(rows_to_delete):
                        tbl = table._tbl
                        tr = tbl.tr_lst[row_idx]
                        tbl.remove(tr)
                self.logger.info(f"Rows containing '{remove_rows_with_string}' removed from tables successfully.")
            except Exception as e:
                self.logger.error(f"Error removing rows with string '{remove_rows_with_string}' from tables: {e}")
                raise Exception(f"Error removing rows with string from tables: {e}")
        elif self.doc_type == 'pdf':
            self.logger.warning("Removing rows with string from tables in PDFs is not supported.")
        else:
            self.logger.error("Unsupported document type for removing rows with string from tables.")
            raise Exception("Unsupported document type for removing rows with string from tables.")

    def remove_empty_columns_in_tables(self):
        """
        Remove empty columns from all tables in the document.

        Parameters:
            None

        Returns:
            None

        Raises:
            Exception:
                If there is an error during processing.

        Upstream functions:
            - `process_document()`

        Downstream functions:
            - None

        Dependencies:
            - Requires the `docx` module.
            - The document must be a DOCX file, and `self.document` must be an instance of `docx.Document`.
            - The parameter `removeEmptyColumns` must be set to `'yes'` in `self.param_manager`.
        """
        remove_empty_columns = self.param_manager.get_parameter('removeEmptyColumns', 'yes')
        if remove_empty_columns.lower() != 'yes':
            self.logger.info("Skipping removing empty columns from tables as per parameters.")
            return

        if self.doc_type == 'docx':
            try:
                for table in self.document.tables:
                    num_cols = len(table.columns)
                    cols_to_delete = []
                    for col_idx in range(num_cols):
                        is_empty = True
                        for row in table.rows:
                            cell = row.cells[col_idx]
                            if cell.text.strip():
                                is_empty = False
                                break
                        if is_empty:
                            cols_to_delete.append(col_idx)
                    # Delete columns from the rightmost to avoid index shifting
                    for col_idx in sorted(cols_to_delete, reverse=True):
                        for row in table.rows:
                            cell = row.cells[col_idx]
                            cell._tc.getparent().remove(cell._tc)
                self.logger.info("Empty columns removed from tables successfully.")
            except Exception as e:
                self.logger.error(f"Error removing empty columns from tables: {e}")
                raise Exception(f"Error removing empty columns from tables: {e}")
        elif self.doc_type == 'pdf':
            self.logger.warning("Removing empty columns from tables in PDFs is not supported.")
        else:
            self.logger.error("Unsupported document type for removing empty columns from tables.")
            raise Exception("Unsupported document type for removing empty columns from tables.")
    

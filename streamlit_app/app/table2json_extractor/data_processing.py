import os
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path

# For Word document parsing
from docx.document import Document as DocxDocument
import docx
from docx.table import _Cell, Table as DocxTable
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

# For PDF parsing
import pdfplumber

# For OCR
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Logging
import logging

# Import custom exceptions
from .exceptions import (
    UnsupportedFileTypeError,
    ParsingError,
    TableExtractionError,
    DocxFileError,
    PDFFileError,
    StructureInterpretationError,
    NestedTableParsingError,
    MergedCellError,
)

# Initialize logger
logger = logging.getLogger(__name__)

class Cell:
    """
    Represents an individual cell within a table.

    Attributes:
        content (Any):
            The text or data within the cell.
        rowspan (int):
            Number of rows this cell spans.
        colspan (int):
            Number of columns this cell spans.
        styles (Dict[str, Any]):
            Styling information (bold, italic, font size, etc.).
        nested_table (Optional[List['Table']]):
            A nested Table object if the cell contains a nested table.
    """

    def __init__(
        self,
        content: Any,
        rowspan: int = 1,
        colspan: int = 1,
        styles: Dict[str, Any] = None,
        nested_table: 'Table' = None
    ):
        self.content = content
        self.rowspan = rowspan
        self.colspan = colspan
        self.styles = styles or {}
        self.nested_table = nested_table

class Table:
    """
    Represents a table extracted from a document.

    Attributes:
        data (List[List[Cell]]): 
            The data of the table as a list of rows, each row being a list of Cell objects.
        position (int): 
            The index position of the table in the document.
        metadata (Dict[str, Any]): 
            Metadata about the table, such as page number, titles, etc.
    """

    def __init__(
        self,
        data: List[List[Cell]],
        position: int = 0,
        metadata: Dict[str, Any] = None
    ):
        self.data = data
        self.position = position
        self.metadata = metadata or {}

class ParsedDocument:
    """
    Represents a parsed document and its content.

    Attributes:
        file_path (str): 
            The file path to the document.
        content (Any): 
            The raw content of the document.
        tables (List[Table]): 
            A list of tables extracted from the document.
        metadata (Dict[str, Any]): 
            Metadata associated with the document.
    """

    def __init__(
        self,
        file_path: str,
        content: Any,
        tables: List[Table],
        metadata: Dict[str, Any] = None
    ):
        self.file_path = file_path
        self.content = content
        self.tables = tables
        self.metadata = metadata or {}

def parse_documents(file_paths: List[str]) -> Tuple[List[ParsedDocument], List[str]]:
    """
    Parses a list of documents and extracts tables from them.

    Parameters:
        file_paths (List[str]):
            List of file paths to the documents to parse.

    Returns:
        Tuple[List[ParsedDocument], List[str]]:
            A tuple containing a list of ParsedDocument objects and a list of any parse errors encountered.
    """

    parsed_documents = []
    parse_errors = []
    for file_path in file_paths:
        file_name = os.path.basename(file_path)
        logger.debug(f"Parsing file: {file_name}")
        try:
            if file_path.lower().endswith('.docx'):
                # Parse DOCX file
                docx_document = docx.Document(file_path)  # Open the document
                tables = parse_docx_tables(docx_document)
                parsed_document = ParsedDocument(file_path=file_path, content=docx_document, tables=tables)
                parsed_documents.append(parsed_document)
                logger.info(f"Parsed {len(tables)} tables from DOCX file '{file_name}'.")
            elif file_path.lower().endswith('.pdf'):
                # Parse PDF file
                with pdfplumber.open(file_path) as pdf:
                    tables = parse_pdf_tables(pdf)
                    parsed_document = ParsedDocument(file_path=file_path, content=pdf, tables=tables)
                    parsed_documents.append(parsed_document)
                    logger.info(f"Parsed {len(tables)} tables from PDF file '{file_name}'.")
            else:
                raise UnsupportedFileTypeError(f"Unsupported file type for file '{file_name}'.")
        except Exception as e:
            logger.exception(f"Error parsing file '{file_name}': {e}")
            parse_errors.append(f"Error parsing file '{file_name}': {e}")
    return parsed_documents, parse_errors


def parse_docx_tables(docx_document: DocxDocument) -> List[Table]:
    """
    Parses tables from a DOCX document.

    Parameters:
        docx_document (DocxDocument):
            The DOCX document object to parse.

    Returns:
        List[Table]:
            A list of Table objects extracted from the document.
    """
    tables = []
    for pos, docx_table in enumerate(docx_document.tables):
        table_data = []
        for row in docx_table.rows:
            row_data = []
            for cell in row.cells:
                # Extract text from cell paragraphs
                cell_text = '\n'.join(paragraph.text for paragraph in cell.paragraphs).strip()
                cell_obj = Cell(content=cell_text)
                row_data.append(cell_obj)
            table_data.append(row_data)
        table = Table(data=table_data, position=pos)
        tables.append(table)
    return tables

def parse_pdf_tables(pdf: pdfplumber.PDF) -> List[Table]:
    """
    Parses tables from a PDF document using pdfplumber.

    Parameters:
        pdf (pdfplumber.PDF):
            The pdfplumber PDF object.

    Returns:
        List[Table]:
            A list of Table objects extracted from the document.

    Raises:
        PDFFileError:
            If an error occurs while parsing the PDF file.
    """
    tables = []
    try:
        for page_num, page in enumerate(pdf.pages):
            page_tables = page.extract_tables()
            for pos, table_data in enumerate(page_tables):
                table_cells = []
                for row in table_data:
                    row_cells = []
                    for cell_text in row:
                        cell_obj = Cell(content=cell_text.strip() if cell_text else '')
                        row_cells.append(cell_obj)
                    table_cells.append(row_cells)
                table = Table(data=table_cells, position=pos)
                tables.append(table)
        return tables
    except Exception as e:
        logger.exception(f"Error parsing PDF: {e}")
        raise PDFFileError(f"Error parsing PDF: {e}") from e

def read_word_document(file_path: str) -> ParsedDocument:
    """
    Reads a Microsoft Word document and prepares it for table extraction.

    Parameters:
        file_path (str): 
            The file path to the Word document.

    Returns:
        ParsedDocument:
            A ParsedDocument object containing the content of the Word file.

    Raises:
        DocxFileError:
            If an error occurs while reading the Word document.
    """
    #try:
    logger.debug(f"Reading Word document: {file_path}")
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    if file_extension == '.doc':
        logger.error("DOC files are not supported. Please convert to DOCX.")
        raise DocxFileError("DOC files are not supported. Please convert to DOCX.")
    # Use the docx.Document function to load the document
    doc = docx.Document(file_path)
    parsed_document = ParsedDocument(file_path=file_path, content=doc)
    return parsed_document
    #except Exception as e:
        #logger.exception(f"Error reading Word document: {file_path}")
        #raise DocxFileError(f"Error reading Word document {file_path}: {str(e)}")

def read_pdf_document(file_path: str) -> ParsedDocument:
    """
    Reads a PDF document and prepares it for table extraction.

    Parameters:
        file_path (str): 
            The file path to the PDF document.

    Returns:
        ParsedDocument:
            A ParsedDocument object containing the content of the PDF file.

    Raises:
        PDFFileError:
            If an error occurs while reading the PDF document.
    """
    try:
        logger.debug(f"Reading PDF document: {file_path}")
        pdf = pdfplumber.open(file_path)
        parsed_document = ParsedDocument(file_path=file_path, content=pdf)
        return parsed_document
    except Exception as e:
        logger.exception(f"Error reading PDF document: {file_path}")
        raise PDFFileError(f"Error reading PDF document {file_path}: {str(e)}")

def extract_raw_tables(document: ParsedDocument) -> List[Table]:
    """
    Extracts raw tables from a single document object.

    Parameters:
        document (ParsedDocument): 
            A ParsedDocument object representing the parsed document content.

    Returns:
        List[Table]:
            A list of Table objects extracted from the document.

    Raises:
        TableExtractionError:
            If an error occurs during table extraction.
    """
    tables = []
    #try:
    if isinstance(document.content, DocxDocument):
        tables = extract_tables_from_word(document.content)
    elif isinstance(document.content, pdfplumber.PDF):
        tables = extract_tables_from_pdf(document.content)
    else:
        logger.error("Unsupported document content type.")
        raise TableExtractionError("Unsupported document content type.")
    return tables
    #except Exception as e:
        #logger.exception(f"Error extracting tables from document {document.file_path}")
        #raise TableExtractionError(f"Error extracting tables from document {document.file_path}: {str(e)}")

def extract_tables_from_word(doc: DocxDocument) -> List[Table]:
    logger.debug("Extracting tables from Word document")
    tables = []
    try:
        all_tables = []
        for block in iter_block_items(doc):
            if isinstance(block, DocxTable):
                all_tables.append(block)

        for idx, word_table in enumerate(all_tables):
            logger.debug(f"Processing table {idx} in Word document")
            table_obj = process_docx_table(word_table, idx)
            tables.append(table_obj)
        return tables
    except Exception as e:
        logger.exception("Error extracting tables from Word document")
        raise TableExtractionError(f"Error extracting tables from Word document: {str(e)}")

def iter_block_items(parent):
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, parent)
        else:
            pass  # Handle other elements if necessary

def process_docx_table(word_table: docx.table.Table, position: int) -> Table:
    try:
        table_data = []
        grid, max_row, max_col = build_table_grid(word_table)
        for r in range(max_row):
            row_data = []
            for c in range(max_col):
                cell = grid[r][c]
                if cell:
                    row_data.append(cell)
                else:
                    empty_cell = Cell(content='')
                    row_data.append(empty_cell)
            table_data.append(row_data)
        table_obj = Table(data=table_data, position=position)
        return table_obj
    except Exception as e:
        logger.exception(f"Error processing table at position {position}")
        raise TableExtractionError(f"Error processing table at position {position}: {str(e)}")

def build_table_grid(word_table: DocxTable) -> Tuple[List[List[Optional[Cell]]], int, int]:
    """
    Builds a grid representing the table, correctly accounting for merged cells.

    Parameters:
        word_table (docx.table.Table): 
            The Word table object.

    Returns:
        Tuple[List[List[Optional[Cell]]], int, int]:
            - The grid as a list of lists containing Cell objects or None.
            - The maximum number of rows.
            - The maximum number of columns.
    """
    grid = {}
    max_row = 0
    max_col = 0

    for row_idx, row in enumerate(word_table.rows):
        col_idx = 0
        for cell in row.cells:
            while (row_idx, col_idx) in grid:
                col_idx += 1

            cell_text = cell.text.strip()
            cell_styles = extract_cell_styles(cell)
            rowspan, colspan = get_cell_span(cell)

            nested_table = None
            if cell.tables:
                nested_table = extract_word_nested_tables(cell)

            cell_obj = Cell(
                content=cell_text,
                rowspan=rowspan,
                colspan=colspan,
                styles=cell_styles,
                nested_table=nested_table,
            )

            # Mark the grid positions occupied by this cell
            for i in range(rowspan):
                for j in range(colspan):
                    r = row_idx + i
                    c = col_idx + j
                    grid[(r, c)] = cell_obj
                    if r + 1 > max_row:
                        max_row = r + 1
                    if c + 1 > max_col:
                        max_col = c + 1
            col_idx += colspan

    # Build the grid
    grid_list = [[None for _ in range(max_col)] for _ in range(max_row)]
    for (r, c), cell_obj in grid.items():
        grid_list[r][c] = cell_obj

    return grid_list, max_row, max_col

def get_cell_span(cell: _Cell) -> Tuple[int, int]:
    """
    Determines the rowspan and colspan of a Word cell.

    Parameters:
        cell (docx.table._Cell): 
            The cell object from which to determine span.

    Returns:
        Tuple[int, int]:
            A tuple containing (rowspan, colspan)
    """
    # Access the underlying XML to find the gridSpan and vMerge attributes
    tc = cell._tc  # The XML element
    grid_span = tc.xpath("./w:tcPr/w:gridSpan")
    v_merge = tc.xpath("./w:tcPr/w:vMerge")

    # Colspan
    if grid_span:
        colspan = int(grid_span[0].get(qn('w:val')))
    else:
        colspan = 1

    # Rowspan
    rowspan = 1
    if v_merge:
        v_merge_val = v_merge[0].get(qn('w:val'))
        if v_merge_val == 'restart' or v_merge_val is None:
            rowspan = count_row_span(cell)
        else:
            rowspan = 0  # This cell is a continuation of a vertically merged cell
    else:
        rowspan = 1

    return rowspan, colspan

def count_row_span(cell: _Cell) -> int:
    """
    Counts how many rows a cell spans in a vertical merge.

    Parameters:
        cell (docx.table._Cell):
            The starting cell of the vertical merge.

    Returns:
        int:
            The number of rows the cell spans.
    """
    rowspan = 1
    current_tc = cell._tc
    current_row = current_tc.getparent().getparent()
    table = current_row.getparent()
    rows = list(table.iterchildren())

    current_row_idx = rows.index(current_row)
    total_rows = len(rows)

    for row_idx in range(current_row_idx + 1, total_rows):
        next_row = rows[row_idx]
        tc_list = next_row.findall('.//w:tc', cell._element.nsmap)
        for next_tc in tc_list:
            next_v_merge = next_tc.xpath("./w:tcPr/w:vMerge")
            if not next_v_merge:
                continue
            next_v_merge_val = next_v_merge[0].get(qn('w:val'))
            if next_v_merge_val == 'continue' or next_v_merge_val is None:
                rowspan += 1
                break  # Found the vertically merged cell in this row
        else:
            break  # No vertically merged cell found in this row

    return rowspan

def extract_cell_styles(cell: _Cell) -> Dict[str, Any]:
    """
    Extracts styles from a Word cell.

    Parameters:
        cell (docx.table._Cell): 
            The cell object from which to extract styles.

    Returns:
        Dict[str, Any]:
            A dictionary of styles applied to the cell.
    """
    styles = {
        'bold': False,
        'italic': False,
        'underline': False,
        'font_size': None,
        'font_name': None,
        'color': None
    }
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                styles['bold'] = True
            if run.italic:
                styles['italic'] = True
            if run.underline:
                styles['underline'] = True
            if run.font.size and (styles['font_size'] is None or run.font.size.pt > styles['font_size']):
                styles['font_size'] = run.font.size.pt
            if run.font.name and styles['font_name'] is None:
                styles['font_name'] = run.font.name
            if run.font.color and run.font.color.rgb and styles['color'] is None:
                styles['color'] = str(run.font.color.rgb)
    return styles

def extract_word_nested_tables(cell: _Cell) -> Optional[List['Table']]:
    """
    Extracts nested tables from a Word document cell.

    Parameters:
        cell (docx.table._Cell): 
            The cell object from which to extract nested tables.

    Returns:
        Optional[List[Table]]:
            A list of Table objects representing the nested tables, if any.
    """
    try:
        nested_tables = []
        for nested_table in cell.tables:
            logger.debug("Extracting nested table from cell")
            grid, max_row, max_col = build_table_grid(nested_table)
            table_data = []
            for r in range(max_row):
                row_data = []
                for c in range(max_col):
                    cell_obj = grid[r][c]
                    if cell_obj:
                        row_data.append(cell_obj)
                    else:
                        empty_cell = Cell(content='')
                        row_data.append(empty_cell)
                table_data.append(row_data)
            nested_table_obj = Table(data=table_data, position=0)
            nested_tables.append(nested_table_obj)
        return nested_tables if nested_tables else None
    except Exception as e:
        logger.exception("Error extracting nested tables from cell")
        raise NestedTableParsingError(f"Error extracting nested tables: {str(e)}")

def extract_tables_from_pdf(pdf: pdfplumber.PDF) -> List[Table]:
    """
    Extracts tables from a PDF document.

    Parameters:
        pdf (pdfplumber.PDF): 
            The PDF document object.

    Returns:
        List[Table]:
            A list of Table objects extracted from the document.

    Raises:
        TableExtractionError:
            If an error occurs during table extraction.
    """
    logger.debug("Extracting tables from PDF document")
    tables = []
    idx = 0
    try:
        for page_number, page in enumerate(pdf.pages, start=1):
            logger.debug(f"Processing page {page_number} in PDF document")
            extracted_tables = page.extract_tables()
            for table in extracted_tables:
                table_data = []
                for row in table:
                    row_data = []
                    for cell_content in row:
                        cell_text = cell_content.strip() if cell_content else ''
                        cell_obj = Cell(content=cell_text)
                        row_data.append(cell_obj)
                    table_data.append(row_data)
                table_obj = Table(
                    data=table_data,
                    position=idx,
                    metadata={'page_number': page_number}
                )
                tables.append(table_obj)
                idx += 1
        return tables
    except Exception as e:
        logger.exception("Error extracting tables from PDF document")
        raise TableExtractionError(f"Error extracting tables from PDF document: {str(e)}")